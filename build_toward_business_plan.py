from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("artifacts/toward_business_plan")
DOCX_PATH = OUT_DIR / "toward.to_business_plan_2026-05-29.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
GRAY = "F2F4F7"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color: str = "D9E2F3") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Title", 22, NAVY, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "toward.to business plan | prepared May 29, 2026"
    footer.runs[0].font.size = Pt(8)
    add_page_number(section.footer.add_paragraph())


def p(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    hdr = tbl.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], BLUE)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if widths:
                set_cell_width(cells[i], widths[i])
    return tbl


def score_table(doc: Document) -> None:
    rows = [
        [
            "Founder/operators and solo businesses",
            "High",
            "25-44; college-educated; AI-comfortable; decisions tied to revenue, hiring, launch, pricing, positioning.",
            "Best beachhead: urgent pain, clear willingness to pay, short feedback loops, and language already present in Toward.",
        ],
        [
            "Consultants, fractional leaders, agencies",
            "High",
            "30-54; professional services; client-facing; decisions must be defensible and shared.",
            "Strong second wedge because PDF/shareable decision briefs can become part of client workflow.",
        ],
        [
            "Career switchers and negotiators",
            "Medium",
            "22-44; knowledge workers; likely to use AI for learning and work; episodic high-stakes decisions.",
            "Good content and template segment, but harder retention unless tied to a sequence of decisions.",
        ],
        [
            "Managers inside larger companies",
            "Medium",
            "30-54; management/professional occupations; higher budget potential but slower procurement.",
            "Attractive later if privacy, admin controls, auditability, and team sharing mature.",
        ],
        [
            "General consumers making purchases/life choices",
            "Low now",
            "Broad age/income mix; intent is fragmented; trust expectations vary widely.",
            "Avoid as launch ICP. Too expensive to message and too close to generic AI assistant positioning.",
        ],
    ]
    table(doc, ["Segment", "Priority", "Likely demographics", "Launch implication"], rows, [1.55, 0.85, 2.2, 2.2])


def gap_table(doc: Document) -> None:
    rows = [
        [
            "1",
            "Launch ICP is too broad",
            "Critical",
            "The site spans career, negotiations, major purchases, startup strategy, hiring, and projects. That makes the product sound useful but not urgent.",
            "Choose one beachhead: founder/operators making business decisions in the next 30 days. Rewrite landing, onboarding examples, advisor defaults, and templates around that wedge.",
        ],
        [
            "2",
            "Production readiness / public source exposure",
            "Critical",
            "The latest site inspected at toward.to serves a Vite-style client app with source routes visible under /src. That is normal in development, but not a launch posture.",
            "Ship a production build, audit source maps and environment exposure, confirm auth/session settings, add monitoring, and run a security/privacy review before paid launch.",
        ],
        [
            "3",
            "No visible monetization architecture",
            "Critical",
            "The public surface sells value but does not show pricing, trial boundaries, paid limits, or a conversion path from decision moment to purchase.",
            "Add a pricing test: free first decision brief, Pro for unlimited decisions/advisor packs/export, Team for shared projects and admin controls.",
        ],
        [
            "4",
            "Trust, liability, and data handling are under-specified",
            "High",
            "Toward touches career, hiring, negotiation, business strategy, and major purchases. Users may upload sensitive files and rely on outputs.",
            "Publish plain-English privacy/TOS, sensitive-data warnings, retention controls, deletion/export guarantees, and disclaimers for legal/financial/medical/hiring advice.",
        ],
        [
            "5",
            "Evidence and auditability gap",
            "High",
            "A 'defensible decision' needs assumptions, evidence, tradeoffs, confidence, and a record of why the recommendation changed.",
            "Add a decision brief format with assumptions register, missing information, alternatives rejected, confidence, and 'what would change this recommendation.'",
        ],
        [
            "6",
            "Onboarding may feel abstract",
            "High",
            "The core product asks for context, constraints, files, advisors, and goals. Without templates, a new user may not know what a good input looks like.",
            "Create ICP-specific decision templates: choose first segment, launch beta, price offer, hire contractor, negotiate contract, cut scope.",
        ],
        [
            "7",
            "Retention loop is not yet obvious",
            "Medium",
            "Decision tools can become one-off utilities unless the product returns to outcomes and learning.",
            "Build follow-up loops: next-action check-ins, decision review, outcome capture, weekly 'open decisions' digest, and saved playbooks.",
        ],
        [
            "8",
            "Distribution surface is thin",
            "Medium",
            "The resource library currently has a small number of decision/productivity articles, so organic acquisition is early.",
            "Build use-case SEO pages and community-led distribution around founder decision moments; collect public examples and case studies.",
        ],
        [
            "9",
            "No launch metrics contract",
            "Medium",
            "The business needs to know whether Toward creates better action, not merely more AI output.",
            "Define activation as first saved/shared decision brief; track time-to-brief, next-action completion, return rate, and paid conversion by decision type.",
        ],
    ]
    table(doc, ["#", "Gap", "Severity", "Why it matters", "Pre-launch requirement"], rows, [0.35, 1.25, 0.75, 2.0, 2.65])
    for row in doc.tables[-1].rows[1:]:
        severity = row.cells[2].text.strip()
        fill = RED if severity == "Critical" else YELLOW if severity == "High" else GREEN
        set_cell_shading(row.cells[2], fill)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    # Cover
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("toward.to Business Plan").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Market focus, demographic rationale, GTM plan, and launch-gap register").font.size = Pt(14)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared for review | Site inspected May 29, 2026").font.size = Pt(10)

    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, "BFBFBF")
    cover_rows = [
        ("Recommended beachhead", "AI-native founder/operators and independent professionals making costly business decisions in the next 30 days."),
        ("Positioning", "A decision lab that turns messy context into a defensible next action, not another task manager or generic chatbot."),
        ("Pre-launch priority", "Narrow the ICP, ship production hardening, add pricing, trust controls, and decision auditability."),
    ]
    for idx, (label, value) in enumerate(cover_rows):
        set_cell_text(tbl.cell(idx, 0), label, bold=True, color="FFFFFF")
        set_cell_shading(tbl.cell(idx, 0), NAVY)
        set_cell_text(tbl.cell(idx, 1), value)
        set_cell_width(tbl.cell(idx, 0), 2.0)
        set_cell_width(tbl.cell(idx, 1), 4.5)

    doc.add_page_break()

    p(doc, "Executive Summary", "Heading 1")
    p(
        doc,
        "Toward is best understood as a decision-intelligence workspace: it captures messy context, lets a user query advisor perspectives, and converts uncertainty into a next best action. The current product language is strong, but it is still too broad for a launch. The market should not hear 'AI productivity app.' The market should hear: 'Make the business decision you can defend.'",
    )
    p(doc, "Recommendation: market first to founder/operators and independent professionals.", "Heading 2")
    bullets(
        doc,
        [
            "Primary ICP: solo founders, small-team operators, consultants, fractional leaders, agency owners, and independent builders who face ambiguous decisions tied to revenue, customers, pricing, launches, hiring, or client delivery.",
            "Why now: AI usage is already mainstream among younger, educated, working adults, but generic chat tools do not preserve context, surface blind spots consistently, or create a durable decision record.",
            "Avoid at launch: broad consumer life decisions, generic productivity, and enterprise management teams. Those markets are larger but noisier, less urgent, or require heavier trust/procurement maturity.",
        ],
    )
    p(doc, "Business Plan Snapshot", "Heading 2")
    table(
        doc,
        ["Area", "Plan"],
        [
            ["Product category", "Decision-intelligence workspace for high-stakes next actions."],
            ["Wedge offer", "One high-quality decision brief in under 20 minutes for a current business decision."],
            ["Audience", "AI-comfortable founder/operators and independent professionals, mostly 25-44, college-educated knowledge workers."],
            ["Model", "Freemium or low-friction trial, then Pro subscription; later team/workspace plan."],
            ["Launch motion", "Use-case landing pages, founder communities, consulting/fractional-leader partnerships, decision-template lead magnets, and example briefs."],
            ["Key risk", "Broad positioning plus weak trust/pricing/auditability can make Toward feel like a nicer wrapper around ChatGPT."],
        ],
        [1.65, 4.95],
    )

    doc.add_page_break()
    p(doc, "What Toward Is Today", "Heading 1")
    p(
        doc,
        "The current public surface positions toward.to around 'Decision Intelligence System,' 'Next Best Action,' and 'Make the higher quality decision, faster.' The app surface includes account login, projects/initiatives, key results, notes/files/constraints/guidance, quick advice, an advisor library, and a blog/resource area.",
    )
    p(doc, "Observed strengths", "Heading 2")
    bullets(
        doc,
        [
            "The product has a coherent decision theme: scattered signals, blind spots, tradeoffs, advisor perspectives, and a next move.",
            "Advisor selection is differentiated enough to feel like a structured thinking system rather than a raw prompt box.",
            "The workspace model can support compounding context across decisions, which is the main strategic advantage over one-off AI chats.",
            "Use cases already include business-relevant moments: startup strategy, hiring, complex projects, prioritization, and launching an invite-only beta.",
        ],
    )
    p(doc, "Observed weaknesses", "Heading 2")
    bullets(
        doc,
        [
            "The use-case spread is too wide for crisp launch messaging: career, negotiations, purchases, startup strategy, hiring, and complex projects all compete for attention.",
            "The site does not yet make the paid conversion path obvious.",
            "The value claim 'defensible decision' needs more explicit evidence handling, assumption tracking, and decision history.",
            "The latest inspected deployment appears to expose development-style source routes. This should be treated as a launch blocker until a production build and security review are complete.",
        ],
    )

    p(doc, "Market Recommendation", "Heading 1")
    p(
        doc,
        "Market Toward first to people who make consequential decisions without a strategy staff. These customers are overloaded with inputs, already comfortable using AI, and willing to pay for clarity because a bad decision has visible cost.",
    )
    p(doc, "Beachhead ICP", "Heading 2")
    table(
        doc,
        ["Attribute", "Definition"],
        [
            ["Who", "Founder/operators, solo business owners, agency owners, consultants, fractional leaders, and independent builders."],
            ["Moment", "A business decision due in the next 30 days: launch, pricing, offer design, segment choice, hiring, contract negotiation, or scope tradeoff."],
            ["Demographics", "Likely 25-44, college-educated, knowledge-work or owner/operator roles, urban/suburban, digitally native, already using AI for work or learning."],
            ["Psychographics", "Ambitious, time-constrained, self-directed, skeptical of generic advice, wants an answer they can explain to a partner, client, investor, or team."],
            ["Budget", "Individual SaaS budget or business card. Test $19-$39/month Pro; $79-$149/month small-team workspace after team features mature."],
            ["Core promise", "Turn a messy decision into a defensible next action and a short decision brief."],
        ],
        [1.45, 5.15],
    )
    p(doc, "Segment Prioritization", "Heading 2")
    score_table(doc)

    doc.add_page_break()
    p(doc, "Demographic Demand Signals", "Heading 1")
    p(
        doc,
        "The best launch market sits where three patterns overlap: high decision load, high AI adoption, and clear economic consequence. The sources below support a younger, educated, professional, founder/operator-oriented wedge rather than a broad consumer launch.",
    )
    table(
        doc,
        ["Signal", "Evidence", "Implication for Toward"],
        [
            [
                "AI adoption is high enough for a specialized AI workflow.",
                "Pew reports that 34% of U.S. adults had used ChatGPT by early 2025, including 58% of adults under 30. Usage is higher among bachelor's and postgraduate degree holders.",
                "Do not spend launch energy teaching the market what AI is. Target users already experimenting with AI and sell structure, judgment, and memory.",
            ],
            [
                "Work use is growing quickly.",
                "Pew reports 28% of employed adults used ChatGPT for work, with higher usage among younger workers and those with higher education.",
                "Toward can compete as the work-grade decision layer above generic AI chats.",
            ],
            [
                "The professional/managerial base is large.",
                "BLS 2025 annual averages list 71.3 million people in management, professional, and related occupations, including 31.0 million in management, business, and financial operations.",
                "The eventual market is large, but launch focus should be narrower than the full knowledge-work base.",
            ],
            [
                "Entrepreneurship remains elevated.",
                "Kauffman reports the U.S. rate of new entrepreneurs increased in 2025 and remained above pre-pandemic levels.",
                "Founder/operator decision moments are abundant and recurring enough for a beachhead.",
            ],
            [
                "Small businesses face uncertain growth conditions.",
                "Federal Reserve Small Business Credit Survey pages report steady employment growth but mixed revenue and future-growth expectations in recent survey cycles.",
                "Small operators have pressure to make better choices with limited resources, which fits Toward's 'next best action' promise.",
            ],
        ],
        [1.45, 2.55, 2.6],
    )
    p(doc, "Practical Demographic Targeting", "Heading 2")
    bullets(
        doc,
        [
            "Age: concentrate acquisition on 25-44. This is old enough to own consequential work decisions and young enough to show strong AI adoption.",
            "Education/occupation: college-educated knowledge workers, founders, consultants, managers, product/marketing/sales operators, and professional-service owners.",
            "Income/business context: users with discretionary tool budgets, business-card purchasing, or direct financial upside from better decisions.",
            "Decision intent: target search and social around urgent decisions, not identity labels alone: 'which customer segment to target,' 'should I launch now,' 'how to price consulting offer,' 'hire contractor or wait,' 'what am I missing in this strategy.'",
        ],
    )

    p(doc, "Positioning", "Heading 1")
    p(doc, "Recommended public positioning", "Heading 2")
    table(
        doc,
        ["Element", "Recommendation"],
        [
            ["Category", "Decision-intelligence workspace."],
            ["One-line promise", "Turn a messy business decision into a defensible next action in 20 minutes."],
            ["Tagline option", "Make the decision you can defend."],
            ["Primary CTA", "Create my decision brief."],
            ["Proof object", "A saved brief with context, assumptions, alternatives, risks, advisor perspectives, and next action."],
            ["Do not lead with", "AI life coach, productivity app, project manager, or generic advisor marketplace."],
        ],
        [1.55, 5.05],
    )
    p(doc, "Messaging Pillars", "Heading 2")
    bullets(
        doc,
        [
            "Clarity under pressure: Toward organizes context when the user is overloaded.",
            "Blind-spot detection: advisor perspectives challenge the user's default framing.",
            "Defensible action: the product explains why a next action is recommended.",
            "Decision memory: each decision becomes a durable record, not a lost chat thread.",
        ],
    )

    doc.add_page_break()
    p(doc, "Competitive Landscape", "Heading 1")
    table(
        doc,
        ["Category", "Examples", "User job", "Toward differentiation"],
        [
            ["Generic AI chat", "ChatGPT, Claude, Gemini", "Ask for advice or brainstorm.", "More structure, saved context, advisor selection, and decision record."],
            ["Productivity/task tools", "Notion, Asana, Todoist, ClickUp", "Track work and tasks.", "Focuses upstream on what decision to make and why, before task execution."],
            ["Coaching/advisory", "Executive coaches, mentors, consultants", "Get perspective and accountability.", "Cheaper, faster, always-on first pass; can feed better inputs into human advisors."],
            ["Strategy templates", "Miro, Coda, worksheets", "Structure thinking manually.", "Automates synthesis and next-action generation from live context."],
            ["Decision frameworks", "Mental-model courses, books", "Learn better thinking.", "Turns frameworks into a working product loop at the moment of decision."],
        ],
        [1.25, 1.4, 1.65, 2.3],
    )
    p(
        doc,
        "The main competitive threat is not a direct clone. It is user inertia: people will try to get the same answer from a generic AI chat. Toward must prove that structure, memory, and evidence produce a materially better decision brief.",
    )

    p(doc, "Business Model", "Heading 1")
    table(
        doc,
        ["Plan", "What to test"],
        [
            ["Free entry", "One complete decision brief with export watermark or limited saved decisions."],
            ["Pro", "$19-$39/month for unlimited active decisions, advisor packs, exports, follow-up reviews, and saved decision history."],
            ["Decision packs", "$9-$19 one-off premium brief for users not ready for a subscription."],
            ["Team", "$79-$149/month for shared decisions, comments, role-based access, admin controls, and team history."],
            ["Services-assisted validation", "Founding customer cohort with white-glove onboarding to learn language, templates, and willingness to pay."],
        ],
        [1.75, 4.85],
    )
    p(doc, "Simple first-year model for planning only", "Heading 2")
    table(
        doc,
        ["Assumption", "Conservative", "Base", "Aggressive"],
        [
            ["Monthly site visitors by month 12", "10,000", "30,000", "75,000"],
            ["Visitor-to-free decision", "3%", "5%", "7%"],
            ["Free-to-paid conversion", "4%", "7%", "10%"],
            ["Average paid ARPU", "$24", "$29", "$39"],
            ["Estimated month-12 MRR", "$2.9k", "$30.5k", "$204.8k"],
        ],
        [2.2, 1.4, 1.4, 1.4],
    )
    p(
        doc,
        "Treat this model as a planning scaffold, not a forecast. The key validation is whether a user with a real decision will complete a brief, trust it, and either pay or return for another decision.",
    )

    doc.add_page_break()
    p(doc, "Go-To-Market Plan", "Heading 1")
    p(doc, "Launch offer", "Heading 2")
    p(
        doc,
        "Offer a 'Founder Decision Brief' workflow: pick a decision type, add context, select advisors, receive a short brief, and schedule a review. The brief is the product's proof object and the conversion artifact.",
    )
    p(doc, "Acquisition channels", "Heading 2")
    table(
        doc,
        ["Channel", "Execution", "Why it fits"],
        [
            ["Founder/operator communities", "Run weekly decision clinics in indie hacker, startup, agency, and solo-business communities.", "Puts Toward in front of users with active decisions and creates examples quickly."],
            ["SEO use-case pages", "Build pages for launch/no-launch, pricing, customer segment, hiring, negotiation, and offer design decisions.", "Captures high-intent queries better than broad productivity content."],
            ["Template lead magnets", "Publish decision brief templates and downloadable checklists.", "Turns abstract decision quality into a tangible artifact."],
            ["Consultant/fractional partnerships", "Give advisors a client-facing brief format they can use before calls.", "Toward can become preparation infrastructure for paid human advice."],
            ["Founder stories", "Publish before/after decision case studies with anonymized context.", "Trust comes from seeing the product reason through real ambiguity."],
        ],
        [1.4, 2.65, 2.55],
    )
    p(doc, "90-Day Validation Roadmap", "Heading 2")
    numbered(
        doc,
        [
            "Weeks 1-2: pick the beachhead, rewrite landing and onboarding examples, define three decision templates, and harden production deployment.",
            "Weeks 3-4: recruit 25 founder/operator users with active decisions; manually review outputs; collect language, objections, and willingness-to-pay data.",
            "Weeks 5-8: launch public Founder Decision Brief offer, pricing test, analytics, and email follow-up loops.",
            "Weeks 9-12: publish case studies, expand templates, add export/share flow, and decide whether to double down on founders or add consultants as the second ICP.",
        ],
    )
    p(doc, "Success metrics", "Heading 2")
    bullets(
        doc,
        [
            "Activation: 40%+ of signed-up users complete and save a first decision brief.",
            "Time-to-value: median time from signup to first brief under 20 minutes.",
            "Trust: 60%+ of completed brief users rate the recommendation as useful or decision-changing.",
            "Behavior: 30%+ complete the recommended next action or return for a decision review.",
            "Revenue: 5-8% free-to-paid conversion from users who complete a brief.",
        ],
    )

    p(doc, "Biggest Business Gaps Before Launch", "Heading 1")
    gap_table(doc)

    doc.add_page_break()
    p(doc, "Launch Readiness Checklist", "Heading 1")
    table(
        doc,
        ["Readiness area", "Minimum bar before launch"],
        [
            ["ICP and message", "One beachhead, one primary decision moment, one clear CTA, and landing copy rewritten around that user."],
            ["Product proof", "Decision brief includes context, assumptions, alternatives, advisor perspectives, risks, and next action."],
            ["Production/security", "Production build, no unintended source/config exposure, monitored auth, data deletion/export, backups, and incident contact."],
            ["Trust/legal", "Privacy policy, terms, sensitive-data guidance, disclaimers, and explicit scope for high-risk advice categories."],
            ["Pricing", "Public pricing or controlled pricing test with free-to-paid conversion path."],
            ["Analytics", "Events for signup, first decision, brief saved/exported, advisor use, rating, return, paid conversion, and churn reason."],
            ["Distribution", "At least six use-case pages, three example briefs, two founder-community campaigns, and one partner channel test."],
        ],
        [1.8, 4.8],
    )
    p(doc, "Decision for founder review", "Heading 2")
    bullets(
        doc,
        [
            "Approve or reject the beachhead: founder/operators and independent professionals.",
            "Pick the first three decision templates: customer segment, launch/no-launch, pricing/offer.",
            "Decide the first monetization test: free first brief into Pro subscription, or paid one-off decision pack.",
            "Assign launch blockers: production deployment, privacy/TOS, pricing, analytics, and decision brief auditability.",
        ],
    )

    p(doc, "Sources and Inspection Notes", "Heading 1")
    table(
        doc,
        ["Source", "Use in plan"],
        [
            ["toward.to public site and client app source inspected May 29, 2026", "Product positioning, routes, use cases, advisor workflow, quick-advice workflow, blog/resource content, and launch-readiness observations."],
            ["Pew Research Center, '34% of U.S. adults have used ChatGPT, about double the share in 2023' (June 25, 2025)", "AI adoption by age, education, and work use."],
            ["U.S. Bureau of Labor Statistics, CPS Annual Averages, Table 11, 2025", "Scale of management, professional, business, and financial occupations."],
            ["Kauffman Indicators of Entrepreneurship, National Report on Early-Stage Entrepreneurship in the United States: 2025", "Entrepreneurship trend and founder/operator market rationale."],
            ["Federal Reserve Small Business Credit Survey report index, 2025-2026", "Small-business performance, owner-demographic reporting, and uncertainty/growth context."],
        ],
        [2.6, 4.0],
    )
    p(
        doc,
        "Important assumption: this plan is based on the public and client-side product surface available during inspection. Server-side implementation, private analytics, customer data, revenue history, and legal documents were not available.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
